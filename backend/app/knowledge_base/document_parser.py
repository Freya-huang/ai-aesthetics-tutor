import os
import re
from typing import Dict, Tuple, List
from docx import Document

from .models import CATEGORY_MAP


def parse_filename_metadata(filename: str) -> Dict[str, str]:
    base_name = os.path.splitext(filename)[0]
    base_name = re.sub(r'\(\d+\)$', '', base_name).strip()
    base_name = re.sub(r'\.md$', '', base_name).strip()
    
    result = {
        "doc_id": "",
        "category": "OTHER",
        "title": base_name,
        "version": "V1"
    }
    
    pattern_with_underscore = r'^([A-Z]+-\d+)[_\s](.+?)[_\s]?(V\d+)?$'
    pattern_with_space = r'^([A-Z]+-\d+)\s+(.+?)$'
    
    match = re.match(pattern_with_underscore, base_name)
    if match:
        result["doc_id"] = match.group(1)
        result["title"] = match.group(2).strip()
        if match.group(3):
            result["version"] = match.group(3)
    else:
        match = re.match(pattern_with_space, base_name)
        if match:
            result["doc_id"] = match.group(1)
            result["title"] = match.group(2).strip()
    
    result["title"] = re.sub(r'[_\s]+V\d+$', '', result["title"]).strip()
    
    if result["doc_id"]:
        prefix = result["doc_id"].split("-")[0]
        if prefix in CATEGORY_MAP:
            result["category"] = prefix
    
    return result


def extract_paragraphs(docx_path: str) -> List[str]:
    doc = Document(docx_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_paragraphs = [p.strip() for p in cell_text.split('\n') if p.strip()]
                    paragraphs.extend(cell_paragraphs)
    
    return paragraphs


def parse_document(docx_path: str) -> Tuple[Dict[str, str], List[str]]:
    filename = os.path.basename(docx_path)
    metadata = parse_filename_metadata(filename)
    metadata["source"] = filename
    paragraphs = extract_paragraphs(docx_path)
    return metadata, paragraphs
